import sys
import os
import google.generativeai as genai
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QComboBox,
    QPushButton, QTextBrowser, QStackedWidget, QMessageBox, QFrame, QFileDialog
)
from PyQt6.QtGui import QFont, QPixmap, QTextDocument
from PyQt6.QtPrintSupport import QPrinter
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from docx import Document

# Developer mode for debugging
DEVELOPER_MODE = True

# Google Gemini API configuration
GEMINI_API_KEY = "YOUR GEMINI API KEY"  # Replace with your actual API key
genai.configure(api_key=GEMINI_API_KEY)

def get_resource_path(relative_path):
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Worker thread to generate content using Google Gemini API
class GoogleWorker(QThread):
    finished = pyqtSignal(str)  # Emits the generated content

    def __init__(self, subject, topic, subtopic, reference, link, mode="summary"):
        super().__init__()
        self.subject = subject
        self.topic = topic
        self.subtopic = subtopic
        self.reference = reference
        self.link = link
        self.mode = mode

    def run(self):
        if self.mode == "summary":
            prompt = (
                f"Summarize the key concepts of the following educational content in 250 words or less. "
                f"Format the answer in HTML using bullet points for key ideas.\n\n"
                f"Subject: {self.subject}\n"
                f"Topic: {self.topic}\n"
                f"Subtopic: {self.subtopic}\n"
                f"Reference: {self.reference}\n"
                f"Resource Link: {self.link}\n"
            )
        elif self.mode == "outline":
            prompt = (
                f"Generate a detailed teaching outline for a class on {self.subject}, focusing on the topic of {self.topic} and subtopic of {self.subtopic}. "
                f"Include sections for introduction, key concepts, learning objectives, suggested activities, assessment ideas, and conclusion. "
                f"Format the outline in HTML with <h2> headings for sections and <ul> for bullet points.\n"
            )
        if DEVELOPER_MODE:
            print(f"{self.mode.capitalize()} Prompt:", prompt)
        try:
            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(prompt)
            content = response.text.strip() if response.text else f"Error generating {self.mode}."
            if DEVELOPER_MODE:
                print(f"Generated {self.mode.capitalize()}:", content)
        except Exception as e:
            content = f"Error generating {self.mode}: {e}"
            if DEVELOPER_MODE:
                print(f"{self.mode.capitalize()} Error:", e)
        self.finished.emit(content)

# SelectionWindow: Left banner + right card for dropdowns
class SelectionWindow(QWidget):
    def __init__(self, resources, switch_window):
        super().__init__()
        self.resources = resources
        self.switch_window = switch_window
        main_layout = QHBoxLayout(self)

        # Left panel: Banner Image
        left_panel = QLabel()
        left_panel.setAlignment(Qt.AlignmentFlag.AlignCenter)
        banner_path = get_resource_path("assets/science_banner.jpeg")
        if not os.path.exists(banner_path):
            banner_path = get_resource_path("assets/science_banner.jpg")
        banner_pixmap = QPixmap(banner_path)
        if banner_pixmap.isNull():
            left_panel.setText("Banner Image Not Found")
            left_panel.setStyleSheet("color: white; font-size: 18px;")
        else:
            left_panel.setPixmap(banner_pixmap.scaled(500, 500, Qt.AspectRatioMode.KeepAspectRatio))

        # Right panel: Card-style selection area
        right_panel = QFrame()
        right_panel.setStyleSheet("""
            background: rgba(255, 255, 255, 0.15);
            border-radius: 15px;
            padding: 20px;
        """)
        right_layout = QVBoxLayout(right_panel)
        right_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        title = QLabel("IKS Resource Portal")
        title.setFont(QFont("Montserrat", 24, QFont.Weight.Bold))
        title.setStyleSheet("color: white;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        right_layout.addWidget(title)

        self.grade_combo = self.create_combobox("Class:", list(self.resources.keys()), right_layout, self.update_subjects)
        self.subject_combo = self.create_combobox("Subject:", [], right_layout, self.update_topics)
        self.topic_combo = self.create_combobox("Topic:", [], right_layout, self.update_subtopics)
        self.subtopic_combo = self.create_combobox("Subtopic:", [], right_layout)

        self.show_button = QPushButton("Show Resources")
        self.show_button.setFont(QFont("Montserrat", 14, QFont.Weight.Bold))
        self.show_button.setStyleSheet("""
            background-color: #1ABC9C; color: white; padding: 10px; border-radius: 10px;
        """)
        self.show_button.clicked.connect(self.show_resources)
        right_layout.addWidget(self.show_button)

        main_layout.addWidget(left_panel, 1)
        main_layout.addWidget(right_panel, 1)

    def create_combobox(self, label_text, values, layout, event_handler=None):
        label = QLabel(label_text)
        label.setFont(QFont("Montserrat", 12, QFont.Weight.Bold))
        label.setStyleSheet("color: white;")
        layout.addWidget(label)
        combo = QComboBox()
        combo.addItems(values)
        combo.setFont(QFont("Montserrat", 11))
        combo.setStyleSheet("""
            QComboBox {
                padding: 6px; border: 2px solid #00A8E8;
                border-radius: 5px; background-color: #FFFFFF;
            }
            QComboBox::drop-down { border: none; }
        """)
        layout.addWidget(combo)
        if event_handler:
            combo.currentIndexChanged.connect(event_handler)
        return combo

    def update_subjects(self):
        grade = self.grade_combo.currentText()
        subjects = list(self.resources.get(grade, {}).keys())
        self.subject_combo.clear()
        self.topic_combo.clear()
        self.subtopic_combo.clear()
        if subjects:
            self.subject_combo.addItems(subjects)
        self.update_topics()

    def update_topics(self):
        grade = self.grade_combo.currentText()
        subject = self.subject_combo.currentText()
        topics = list(self.resources.get(grade, {}).get(subject, {}).keys())
        self.topic_combo.clear()
        self.subtopic_combo.clear()
        if topics:
            self.topic_combo.addItems(topics)
        self.update_subtopics()

    def update_subtopics(self):
        grade = self.grade_combo.currentText()
        subject = self.subject_combo.currentText()
        topic = self.topic_combo.currentText()
        subtopics = list(self.resources.get(grade, {}).get(subject, {}).get(topic, {}).keys())
        self.subtopic_combo.clear()
        if subtopics:
            self.subtopic_combo.addItems(subtopics)

    def show_resources(self):
        grade = self.grade_combo.currentText()
        subject = self.subject_combo.currentText()
        topic = self.topic_combo.currentText()
        subtopic = self.subtopic_combo.currentText()
        self.switch_window(grade, subject, topic, subtopic)

# ResourceWindow: Displays resources, summary, and teaching outline button
class ResourceWindow(QWidget):
    def __init__(self, resources, switch_back):
        super().__init__()
        self.resources = resources
        self.switch_back = switch_back
        self.summary_worker = None
        self.outline_worker = None

        main_layout = QVBoxLayout(self)
        main_layout.setAlignment(Qt.AlignmentFlag.AlignCenter)

        title = QLabel("Resources & Links")
        title.setFont(QFont("Montserrat", 24, QFont.Weight.Bold))
        title.setStyleSheet("color: white;")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        main_layout.addWidget(title)

        # Card for summary and outline button
        card = QFrame()
        card.setStyleSheet("""
            background: rgba(255,255,255,0.85);
            border-radius: 15px;
            padding: 15px;
        """)
        card_layout = QVBoxLayout(card)
        self.summary_label = QLabel("Summary will appear here.")
        self.summary_label.setFont(QFont("Montserrat", 12))
        self.summary_label.setStyleSheet("color: #2C3E50;")
        self.summary_label.setWordWrap(True)
        card_layout.addWidget(self.summary_label)

        self.generate_outline_button = QPushButton("Generate Teaching Outline")
        self.generate_outline_button.setFont(QFont("Montserrat", 12, QFont.Weight.Bold))
        self.generate_outline_button.setStyleSheet("""
            background-color: #1ABC9C; color: white; padding: 8px; border-radius: 5px;
        """)
        self.generate_outline_button.clicked.connect(self.start_outline_generation)
        card_layout.addWidget(self.generate_outline_button)

        main_layout.addWidget(card)

        # Resource text area
        self.content_area = QTextBrowser()
        self.content_area.setFont(QFont("Montserrat", 12))
        self.content_area.setStyleSheet("""
            background-color: rgba(255,255,255,0.9);
            padding: 10px; border-radius: 8px; color: #2C3E50;
        """)
        self.content_area.setOpenExternalLinks(True)
        main_layout.addWidget(self.content_area)

        self.back_button = QPushButton("Back to Selection")
        self.back_button.setFont(QFont("Montserrat", 14, QFont.Weight.Bold))
        self.back_button.setStyleSheet("""
            background-color: #FF5733; color: white; padding: 10px; border-radius: 8px;
        """)
        self.back_button.clicked.connect(self.switch_back)
        main_layout.addWidget(self.back_button)

    def display_resources(self, grade, subject, topic, subtopic):
        self.grade = grade  # Store for outline generation
        self.subject = subject
        self.topic = topic
        self.subtopic = subtopic
        data = self.resources.get(grade, {}).get(subject, {}).get(topic, {}).get(subtopic, {})
        if data:
            reference = data["Reference"]
            links = data["Link"].split()
            result_html = (
                f"<div style='font-family: Montserrat; font-size: 14px; color: #2C3E50;'>"
                f"<p><b>Reference:</b> {reference}</p>"
                f"<p><b>Links:</b></p>"
            )
            valid_link = None
            for link in links:
                if link.startswith("http"):
                    result_html += f"<p><a href='{link}' style='color: #1ABC9C;'>{link}</a></p>"
                    if valid_link is None:
                        valid_link = link
            result_html += "</div>"
            self.content_area.setHtml(result_html)
            if valid_link:
                if self.summary_worker is not None:
                    self.summary_worker.terminate()
                self.summary_worker = GoogleWorker(subject, topic, subtopic, reference, valid_link, mode="summary")
                self.summary_worker.finished.connect(self.update_summary)
                self.summary_worker.start()
            else:
                self.summary_label.setText("No valid link available for summary.")
        else:
            self.content_area.setHtml("<div>No resources found.</div>")
            self.summary_label.setText("")

    def update_summary(self, summary):
        self.summary_label.setText(f"<div style='font-family: Montserrat; font-size: 14px; color: #2C3E50;'><b>Summary:</b><br>{summary}</div>")

    def start_outline_generation(self):
        data = self.resources.get(self.grade, {}).get(self.subject, {}).get(self.topic, {}).get(self.subtopic, {})
        if data:
            reference = data["Reference"]
            links = data["Link"].split()
            valid_link = next((link for link in links if link.startswith("http")), None)
            if valid_link:
                if self.outline_worker is not None:
                    self.outline_worker.terminate()
                self.outline_worker = GoogleWorker(self.subject, self.topic, self.subtopic, reference, valid_link, mode="outline")
                self.outline_worker.started.connect(lambda: self.generate_outline_button.setText("Generating..."))
                self.outline_worker.finished.connect(self.on_outline_generated)
                self.outline_worker.start()
            else:
                QMessageBox.warning(self, "Error", "No valid link available.")
        else:
            QMessageBox.warning(self, "Error", "No resources found.")

    def on_outline_generated(self, outline_html):
        self.generate_outline_button.setText("Generate Teaching Outline")
        if outline_html.startswith("Error"):
            QMessageBox.warning(self, "Error", outline_html)
            return
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Teaching Outline", "", "PDF Files (*.pdf)")
        if file_path:
            self.generate_pdf(outline_html, file_path)
            QMessageBox.information(self, "Success", "Teaching outline saved successfully.")

    def generate_pdf(self, html, file_path):
        document = QTextDocument()
        document.setHtml(html)
        printer = QPrinter()
        printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
        printer.setOutputFileName(file_path)
        document.print(printer)

# MainApp: Manages window transitions
class MainApp(QMainWindow):
    def __init__(self, base_folder):
        super().__init__()
        self.setWindowTitle("IKS Resource Portal")
        self.setGeometry(100, 100, 1200, 800)
        self.base_folder = base_folder
        self.resources = self.fetch_data()
        self.stack = QStackedWidget()
        self.selection_screen = SelectionWindow(self.resources, self.switch_to_resources)
        self.resource_screen = ResourceWindow(self.resources, self.switch_to_selection)
        self.stack.addWidget(self.selection_screen)
        self.stack.addWidget(self.resource_screen)
        self.setCentralWidget(self.stack)

    def fetch_data(self):
        resources = {}
        for grade_folder in os.listdir(self.base_folder):
            grade_path = os.path.join(self.base_folder, grade_folder)
            if os.path.isdir(grade_path):
                grade_name = grade_folder.rstrip("th")
                resources[grade_name] = {}
                for subject_folder in os.listdir(grade_path):
                    subject_path = os.path.join(grade_path, subject_folder)
                    if os.path.isdir(subject_path):
                        mapping_file = self.find_mapping_file(subject_path)
                        if mapping_file:
                            resources[grade_name][subject_folder.capitalize()] = self.parse_mapping_file(mapping_file)
        return resources

    def find_mapping_file(self, folder_path):
        for file_name in os.listdir(folder_path):
            if file_name.endswith("Mapping.docx"):
                return os.path.join(folder_path, file_name)
        return None

    def parse_mapping_file(self, file_path):
        data = {}
        doc = Document(file_path)
        current_topic = None
        for table in doc.tables:
            for row in table.rows[1:]:
                topic = row.cells[0].text.strip()
                subtopic = row.cells[1].text.strip()
                reference = row.cells[2].text.strip()
                link = row.cells[3].text.strip()
                if topic:
                    current_topic = topic
                    data[current_topic] = {}
                if current_topic and subtopic:
                    data[current_topic][subtopic] = {"Reference": reference, "Link": link}
        return data

    def switch_to_resources(self, grade, subject, topic, subtopic):
        self.resource_screen.display_resources(grade, subject, topic, subtopic)
        self.stack.setCurrentWidget(self.resource_screen)

    def switch_to_selection(self):
        self.stack.setCurrentWidget(self.selection_screen)

if __name__ == "__main__":
    BASE_FOLDER = get_resource_path("IKS_PAPER")
    app = QApplication(sys.argv)
    app.setStyleSheet("""
        QWidget {
            background: qlineargradient(x1:0, y1:0, x2:0, y2:1, stop:0 #2C3E50, stop:1 #4CA1AF);
        }
    """)
    window = MainApp(BASE_FOLDER)
    window.show()
    sys.exit(app.exec())
